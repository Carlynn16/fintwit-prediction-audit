"""(Re)build report/report.docx from scratch with all text, figures, and tables.

Rebuilds the document completely on every run (deletes existing docx first).
Run scripts/make_figures.py beforehand to generate figure PNGs, OR let main()
call the figure-generators automatically.

Usage:
    python scripts/build_report.py
"""

from __future__ import annotations

import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).parent.parent))
# Also put scripts/ on the path so we can import make_figures as a sibling module
sys.path.insert(0, str(pathlib.Path(__file__).parent))

import numpy as np
import pandas as pd

from src.report import (
    add_body, add_figure, add_heading1, add_heading2, add_table, new_document,
    write_cover_page,
)
from src.stats_rq1 import compute_rq1, segment_stats, wilson_ci
from src.stats_rq2 import compute_rq2, MIN_N
from src.stats_rq3 import calibration_table
from src.stats_regime import compute_regime

REPORT_PATH = pathlib.Path(__file__).parent.parent / "report" / "report.docx"
REPORT_PATH.parent.mkdir(exist_ok=True)

FIGURES_DIR = pathlib.Path(__file__).parent.parent / "figures"
FIGURES_DIR.mkdir(exist_ok=True)

DATA_DIR  = pathlib.Path(__file__).parent.parent / "data"
INPUT_CSV = DATA_DIR / "tweets_validated.csv"

# ── figure paths ─────────────────────────────────────────────────────────────
FIG_TWEETS_MONTH   = FIGURES_DIR / "data_tweets_per_month.png"
FIG_PREDS_ACCOUNT  = FIGURES_DIR / "data_predictions_per_account.png"
FIG_TOP_TICKERS    = FIGURES_DIR / "data_top_tickers.png"
FIG_ACCURACY_SEG   = FIGURES_DIR / "rq1_accuracy_by_segment.png"
FIG_ACCURACY_YEAR  = FIGURES_DIR / "rq1_accuracy_by_year.png"
FIG_EXCESS_RET     = FIGURES_DIR / "rq1_excess_return_hist.png"
FIG_CATERPILLAR    = FIGURES_DIR / "rq2_skill_caterpillar.png"
FIG_MULTTEST       = FIGURES_DIR / "rq2_multiple_testing_bar.png"
FIG_CALIB_CURVE    = FIGURES_DIR / "rq3_calibration_curve.png"
FIG_FEAT_IMPORT    = FIGURES_DIR / "rq3_feature_importance.png"
FIG_ROC            = FIGURES_DIR / "rq3_roc_curves.png"
FIG_CONF_HIST      = FIGURES_DIR / "rq3_confidence_hist.png"
FIG_REGIME         = FIGURES_DIR / "rq4_regime_scatter.png"


# ── table-data helpers ────────────────────────────────────────────────────────

def _pstar(p: float) -> str:
    if p < 0.001: return "***"
    if p < 0.01:  return "**"
    if p < 0.05:  return "*"
    return ""


def _segment_table_rows(rq1: dict) -> list[list[str]]:
    """Build rows for Table 6.1 from compute_rq1() output."""
    rows: list[list[str]] = []

    def _row(label, n, k_c, acc_c, ci_c_lo, ci_c_hi, k_b, acc_b, ci_b_lo, ci_b_hi):
        return [
            label,
            f"{int(n):,}",
            f"{acc_c*100:.1f}%",
            f"[{ci_c_lo*100:.1f}%, {ci_c_hi*100:.1f}%]",
            f"{acc_b*100:.1f}%",
            f"[{ci_b_lo*100:.1f}%, {ci_b_hi*100:.1f}%]",
        ]

    # Overall
    oc = rq1["overall"]["correct"]
    ob = rq1["overall"]["beats"]
    n  = rq1["n_validated"]
    rows.append(_row(
        "Overall (all validated)", n,
        int(oc["k"]), oc["p_hat"], oc["ci_low"], oc["ci_high"],
        int(ob["k"]), ob["p_hat"], ob["ci_low"], ob["ci_high"],
    ))

    segs = rq1["segments"]
    for group_label, seg_key, order in [
        ("Sentiment", "sentiment",    ["bearish", "bullish"]),
        ("Horizon",   "time_horizon", ["short_term", "medium_term", "long_term", "unknown"]),
        ("Trade type","trade_type",   ["analysis", "news", "general_discussion", "trade_suggestion"]),
    ]:
        seg = segs[seg_key]
        for name in order:
            if name not in seg.index:
                continue
            r = seg.loc[name]
            rows.append(_row(
                f"  {group_label}: {name}",
                int(r["n"]),
                int(r["k_correct"]), r["acc_correct"], r["ci_correct_low"], r["ci_correct_high"],
                int(r["k_beats"]),   r["acc_beats"],   r["ci_beats_low"],   r["ci_beats_high"],
            ))

    return rows


def _account_table_rows(rq2: dict) -> list[list[str]]:
    """Build rows for Table 7.1 from compute_rq2() output, sorted by shrunk_beats desc."""
    acc = rq2["account_df"]  # already sorted
    rows = []
    for author, r in acc.iterrows():
        star = _pstar(float(r["pval_beats"]))
        above = " ↑" if r["credible_above_50"] else ""
        rows.append([
            str(author),
            str(int(r["n"])),
            f"{r['beats_rate']*100:.1f}%",
            f"{r['shrunk_beats']*100:.1f}%",
            f"[{r['credible_low']*100:.1f}%, {r['credible_high']*100:.1f}%]{above}",
            f"{r['pval_beats']:.4f}{star}",
            f"{r['pval_beats_bh']:.4f}",
            f"{r['pval_beats_bonf']:.4f}",
        ])
    return rows


def _calib_table_rows(calib: pd.DataFrame) -> list[list[str]]:
    rows = []
    for _, r in calib.iterrows():
        rows.append([
            f"[{r['bin_lower']:.0f}, {r['bin_upper']:.0f})",
            f"{int(r['n_predictions']):,}",
            f"{r['actual_accuracy']*100:.1f}%",
            f"[{r['ci_low']*100:.1f}%, {r['ci_high']*100:.1f}%]",
        ])
    return rows


# ── write_phase functions ─────────────────────────────────────────────────────

def write_phase0(doc) -> None:
    # ── 1. Introduction ────────────────────────────────────────────────────────
    add_heading1(doc, "1. Introduction")
    add_body(doc, (
        "Financial social media — 'FinTwit' — hosts a large and vocal community of "
        "traders and investors who publish directional stock predictions publicly and "
        "in real time. These accounts attract substantial followings and, implicitly, "
        "influence retail investment behaviour. Yet there exists almost no rigorous, "
        "large-scale audit of whether these predictions have any actual merit. This "
        "project fills that gap."
    ))
    add_body(doc, (
        "Using approximately 18,000 tweets from 51 financial-influencer accounts on "
        "X (formerly Twitter), spanning January 2021 to February 2025, we conduct a "
        "reproducible statistical audit that answers three pre-registered research "
        "questions:"
    ))

    add_heading2(doc, "1.1 Research Questions")
    add_body(doc, (
        "RQ1 — Accuracy & benchmark. How accurate are FinTwit predictions "
        "directionally, and do they beat a passive market benchmark (buy-and-hold of "
        "the named stock, and the broad market via SPY) over the matched prediction "
        "horizon?"
    ))
    add_body(doc, (
        "RQ2 — Skill vs luck. Across all 51 accounts, after correcting for multiple "
        "comparisons, is any individual account's track record statistically "
        "distinguishable from chance? Model true per-account skill with an "
        "empirical-Bayes / hierarchical shrinkage estimator that accounts for "
        "small-sample noise."
    ))
    add_body(doc, (
        "RQ3 — Signal value. Do tweet-level signals — sentiment, trade type, "
        "horizon, engagement, and the LLM's own self-reported confidence — carry any "
        "information about whether a prediction comes true? Includes (a) a formal "
        "calibration analysis of the LLM confidence score (a preliminary check "
        "already shows it is essentially uncalibrated, r ≈ 0), and (b) a predictive "
        "model of prediction_correct."
    ))

    add_heading2(doc, "1.2 Provenance")
    add_body(doc, (
        "This is an anonymised personal rebuild of a former paid consulting "
        "engagement, presented as an independent portfolio project. No original "
        "client is named or referenced anywhere in this analysis."
    ))

    # ── 2. Data ────────────────────────────────────────────────────────────────
    add_heading1(doc, "2. Data")

    add_heading2(doc, "2.1 Overview")
    add_body(doc, (
        "The dataset comprises 18,071 tweets posted by 51 financial-influencer "
        "accounts on X between 2021-01-07 and 2025-02-20. The raw tweet text and "
        "all platform metadata were collected prior to this analysis; no Twitter/X "
        "API access is required to reproduce the results. All 51 account handles "
        "have been pseudonymised to account_01 … account_51, assigned "
        "deterministically by alphabetical sort order of the original handles. The "
        "handle-to-pseudonym mapping is stored in data/account_mapping.csv, which "
        "is excluded from version control."
    ))

    add_heading2(doc, "2.2 Derivation of tweets_raw.csv")
    add_body(doc, (
        "The working dataset data/tweets_raw.csv was derived from the source export "
        "by (i) retaining only genuinely raw columns (tweet identifiers, author, "
        "tweet text, timestamps, reply structure, engagement counts, and account "
        "metadata), (ii) dropping all LLM-derived and validation columns that will "
        "be regenerated in later phases (tickers_mentioned, time_horizon, "
        "trade_type, sentiment, confidence, price_change_pct, prediction_correct, "
        "actual_return, and related fields), (iii) parsing types (tz-aware "
        "datetimes, nullable integers, booleans), and (iv) replacing all real "
        "account handles with their pseudonyms."
    ))

    add_heading2(doc, "2.3 Columns")
    add_body(doc, (
        "tweets_raw.csv contains 17 columns. Full definitions, dtypes, and notes "
        "are provided in data/DATA_DICTIONARY.md. In brief: tweet_id, "
        "conversation_id, and reply_to_tweet_id are opaque string identifiers; "
        "tweet_type records the thread role (parent or reply); author and "
        "reply_to_user are pseudonymised handles; text is the raw tweet body and "
        "the sole NLP input for all subsequent phases; created_at is a tz-aware "
        "UTC timestamp and created_date is the corresponding calendar date; likes, "
        "retweets, replies_count, and views are nullable integer engagement counts; "
        "author_followers, author_following are nullable integer account-size "
        "metrics; author_verified and author_blue_verified are booleans."
    ))

    add_heading2(doc, "2.4 Data-quality caveats")
    add_body(doc, (
        "ID precision: the source export stored tweet_id, conversation_id, and "
        "reply_to_tweet_id as float64 (scientific notation, e.g. 1.8924e+18). "
        "Float representation of 64-bit tweet IDs loses the least-significant "
        "digits, so these identifiers are treated as opaque labels only and are "
        "never used as exact join keys against external datasets."
    ))
    add_body(doc, (
        "Missing engagement data: engagement and account-size columns contain "
        "missing values, particularly for older tweets and reply-type tweets. "
        "Nullable Int64 dtype is used throughout to preserve the distinction "
        "between a genuine zero and a missing observation."
    ))

    add_figure(doc, FIG_TWEETS_MONTH,
               "Figure 2.1. Monthly tweet volume across all 51 accounts (2021–2025). "
               "Dashed line shows the overall monthly mean.")
    add_figure(doc, FIG_PREDS_ACCOUNT,
               "Figure 2.2. Number of validated predictions per account (all 51 accounts). "
               "Dashed red line marks the n = 30 minimum threshold used in RQ2.")
    add_figure(doc, FIG_TOP_TICKERS,
               "Figure 2.3. Top 15 most frequently predicted tickers in the validated dataset.")


def write_phase1(doc) -> None:
    # ── 3. Ticker extraction ───────────────────────────────────────────────────
    add_heading1(doc, "3. Ticker Extraction")
    add_body(doc, (
        "Before any directional prediction can be validated against market data, "
        "we must identify which stock(s) each tweet is referring to. We use a "
        "two-stage pipeline that prioritises speed and cost efficiency: a regex "
        "baseline captures the vast majority of cases, and an LLM fallback handles "
        "the remainder."
    ))

    add_heading2(doc, "3.1 Regex Baseline")
    add_body(doc, (
        "The primary extraction method is a regular expression that matches "
        "cashtag patterns of the form $TICKER, where TICKER is between one and "
        "five ASCII letters. The pattern is case-insensitive and requires that no "
        "additional letters immediately follow the match (word boundary), ensuring "
        "that phrases like '$AAPLisgreat' are not falsely matched."
    ))
    add_body(doc, (
        "A stoplist removes tokens that match the pattern but are not stock "
        "tickers: fiat-currency abbreviations (USD, EUR, GBP, JPY, CAD, AUD, CHF, "
        "CNY, and others), and common financial abbreviations that appear as "
        "cashtags (IPO, ETF, CEO, AI, etc.). Dollar-prefixed numbers such as $100 "
        "are excluded automatically because the pattern requires letters. All "
        "surviving matches are upper-cased and de-duplicated within each tweet."
    ))

    add_heading2(doc, "3.2 LLM Fallback")
    add_body(doc, (
        "The LLM fallback is invoked only for tweets in which the regex finds no "
        "cashtag — roughly 27.6% of the corpus. This design keeps API costs "
        "proportional to the hardest cases rather than the full dataset. The "
        "fallback uses gpt-4o-mini with a zero-temperature structured-JSON prompt "
        "that instructs the model to resolve unambiguous company-name mentions to "
        "their primary US ticker (e.g. 'Tesla' to TSLA), return [] when no "
        "specific stock is discussed, and never invent a ticker it is uncertain "
        "about."
    ))
    add_body(doc, (
        "Results are cached on disk (cache/llm_tickers/) keyed by the SHA-256 "
        "hash of the tweet text. A re-run on the same data incurs zero additional "
        "API cost. The OpenAI API key is read from the OPENAI_API_KEY environment "
        "variable; extraction degrades gracefully to regex-only when the key is "
        "absent. Rate-limit and transient API errors are retried with exponential "
        "back-off (up to four attempts)."
    ))

    add_heading2(doc, "3.3 Output Columns")
    add_body(doc, (
        "The extraction step adds two columns to the working dataset "
        "(data/tweets_with_tickers.csv): tickers_mentioned, a Python list of "
        "deduplicated, upper-cased ticker strings (may be empty); and has_ticker, "
        "a boolean that is True if and only if tickers_mentioned is non-empty. "
        "No filtering for tradeable / exchange-listed status is applied at this "
        "stage — that validation occurs in Phase 3 via yfinance."
    ))

    add_heading2(doc, "3.4 Coverage")
    add_body(doc, (
        "Running the full two-stage pipeline over all 18,071 tweets produced the "
        "following coverage:"
    ))
    add_body(doc, (
        "Regex hit: 13,089 tweets (72.4%). "
        "LLM-only hit: 950 tweets (5.3%). "
        "No ticker found: 4,032 tweets (22.3%). "
        "At least one ticker (regex + LLM combined): 14,039 tweets (77.7%)."
    ))
    add_body(doc, (
        "The LLM fallback added 950 tweets that the regex missed — recovering "
        "5.3 percentage points of coverage by resolving company-name mentions "
        "that lacked a cashtag (e.g. 'Tesla' to TSLA). Roughly 4,032 tweets "
        "(22.3%) yielded no ticker from either method, consistent with the "
        "expectation that a share of the corpus consists of market commentary, "
        "retweets, or general observations not tied to a specific stock. "
        "All 4,982 no-cashtag tweets were sent through the LLM; results are "
        "cached on disk (cache/llm_tickers/, 4,705 cache files) so re-runs "
        "make zero additional API calls."
    ))

    add_heading2(doc, "3.5 Evaluation Methodology and Results")
    add_body(doc, (
        "To assess extraction quality, we drew a stratified random sample of "
        "100 tweets (random seed 42) covering both cashtag-containing and "
        "non-cashtag tweets. Each tweet's text was read and the correct tickers "
        "entered by hand, blind to the automated outputs stored in "
        "data/ticker_eval_sample.csv. Once labelled, scripts/eval_tickers.py "
        "computed precision, recall, and F1 at the ticker level, separately for "
        "the regex-only and combined (regex + LLM) methods."
    ))
    add_body(doc, "Table 3.1 — Ticker extraction performance (n = 100 hand-labelled tweets)")
    add_table(
        doc,
        headers=["Method", "TP", "FP", "FN", "Precision", "Recall", "F1"],
        rows=[
            ["Regex only",  "114", "3", "8", "0.974", "0.934", "0.954"],
            ["Regex + LLM", "117", "4", "5", "0.967", "0.959", "0.963"],
        ],
        col_widths_cm=[3.8, 1.2, 1.2, 1.2, 2.2, 2.2, 2.0],
    )
    add_body(doc, (
        "The regex baseline achieves high precision (0.974) — virtually every "
        "cashtag it extracts is a genuine ticker — but misses 8 tickers that "
        "appear as company-name mentions without a cashtag (recall 0.934). "
        "Adding the LLM fallback recovers 3 of those 8 misses, raising recall "
        "to 0.959 at the cost of just one additional false positive (FP 3 → 4), "
        "yielding a net F1 improvement from 0.954 to 0.963. "
        "Both methods perform well; the LLM contributes a modest but consistent "
        "recall gain for company-name mentions."
    ))


def write_phase2(doc) -> None:
    add_heading1(doc, "4. Classification Methodology")
    add_body(doc, (
        "Each of the 14,039 ticker-bearing tweets was classified by gpt-4o-mini "
        "into four dimensions that define the nature of the prediction. Tweets "
        "without a ticker (4,032 rows) are left unclassified and receive null "
        "values for all classification columns."
    ))

    add_heading2(doc, "4.1 Classification Schema")
    add_body(doc, (
        "The model was asked to return a strict JSON object with exactly four fields:"
    ))
    add_body(doc, (
        "sentiment — the tweet's directional view on the named stock(s): "
        "\"bullish\" (positive or long bias), \"bearish\" (negative or short bias), "
        "or \"neutral\" (no clear directional view)."
    ))
    add_body(doc, (
        "time_horizon — the implied prediction window: \"short_term\" (days to "
        "approximately one month), \"medium_term\" (approximately one to six months), "
        "\"long_term\" (more than six months), or \"unknown\" when no horizon is implied."
    ))
    add_body(doc, (
        "trade_type — the nature of the tweet: \"trade_suggestion\" (an explicit "
        "buy, sell, entry, or exit recommendation), \"analysis\" (a reasoned opinion "
        "or chart reading without an explicit trade call), \"news\" (relaying factual "
        "information such as earnings or macro data), or \"general_discussion\" "
        "(everything else)."
    ))
    add_body(doc, (
        "confidence — an integer from 0 to 100 representing how assertive or "
        "certain the tweet sounds, where 0 is very tentative and 100 is extremely "
        "certain. This score reflects the linguistic register of the tweet, not a "
        "calibrated probability of correctness; its calibration against actual "
        "prediction outcomes is assessed separately in Phase 6 (RQ3)."
    ))

    add_heading2(doc, "4.2 Implementation and Caching")
    add_body(doc, (
        "Classification calls used gpt-4o-mini at temperature=0 with "
        "response_format=json_object to enforce structured output. The API key is "
        "read from the OPENAI_API_KEY environment variable. Rate-limit and transient "
        "API errors are retried with exponential back-off (up to four attempts). "
        "On any failure — including malformed or truncated JSON — the row falls back "
        "to safe defaults (sentiment: neutral, time_horizon: unknown, trade_type: "
        "general_discussion, confidence: 0) and the event is logged; the pipeline "
        "never crashes mid-run."
    ))
    add_body(doc, (
        "Results are cached on disk (cache/llm_classify/) keyed by the SHA-256 "
        "hash of the tweet text. Re-running the classification script on the same "
        "corpus makes zero additional API calls; all 14,039 results are served "
        "from the cache."
    ))

    add_heading2(doc, "4.3 Results and Distributions")
    add_body(doc, (
        "All 14,039 ticker-bearing tweets were classified successfully. "
        "13,730 calls were made to the API; 309 were served from the on-disk cache "
        "(duplicate tweet texts). The resulting distributions are as follows."
    ))
    add_body(doc, (
        "Sentiment: bullish 8,847 (63.0%), bearish 2,951 (21.0%), "
        "neutral 2,241 (16.0%). The strong bullish skew is consistent with the "
        "promotional nature of financial-influencer content on X."
    ))
    add_body(doc, (
        "Time horizon: short_term 6,718 (47.9%), unknown 3,793 (27.0%), "
        "medium_term 2,105 (15.0%), long_term 1,423 (10.1%). Nearly half of "
        "all predictions are framed as short-term calls, and more than a quarter "
        "carry no explicit time frame."
    ))
    add_body(doc, (
        "Trade type: analysis 6,307 (44.9%), trade_suggestion 3,263 (23.2%), "
        "general_discussion 2,899 (20.6%), news 1,570 (11.2%). The largest "
        "category is reasoned opinion without an explicit trade call; explicit "
        "buy/sell recommendations account for roughly one in four classified tweets."
    ))
    add_body(doc, (
        "Confidence scores across all 14,039 classified tweets: mean 74.2, "
        "median 75, min 10, max 100. The high median confidence reflects the "
        "assertive register typical of FinTwit content. Whether this expressed "
        "confidence bears any relationship to actual prediction accuracy is "
        "assessed in Phase 6 (RQ3 calibration analysis); preliminary findings "
        "suggest it does not."
    ))

    add_heading2(doc, "4.4 Worked Example")
    add_body(doc, (
        "Table 4.1 shows three representative tweets selected to illustrate the "
        "classification schema and downstream outcomes. Each example contains no "
        "account handles (anonymisation preserved). Tweet text is truncated to "
        "approximately 90 characters; all @-mentions are redacted to @user."
    ))
    _worked_example_table(doc)


def _worked_example_table(doc) -> None:
    headers = ["Tweet text (~90 chars)", "Ticker", "Sentiment",
               "Horizon", "Trade type", "Conf.", "Outcome"]
    rows = [
        [
            "$NVDA to test 110–115",
            "NVDA", "bullish", "short_term", "trade_suggestion", "75", "Correct",
        ],
        [
            "Someone bought $700,000 worth of $TSLA 305 puts that expire next week "
            "— calling for a test of $300",
            "TSLA", "bearish", "short_term", "analysis", "75", "Correct",
        ],
        [
            "Starting to think AI lending and insurance might be a really nice trade "
            "in the coming months. Everyone hates legacy insurance...",
            "LMND", "bullish", "medium_term", "trade_suggestion", "75", "Wrong",
        ],
    ]
    add_body(doc, "Table 4.1. Classification worked example (three real tweets, text redacted/truncated).")
    add_table(
        doc, headers, rows,
        col_widths_cm=[5.8, 1.1, 1.5, 2.0, 2.2, 1.0, 1.3],
        font_size_pt=9,
    )


def write_phase3(doc) -> None:
    add_heading1(doc, "5. Market Validation")
    add_body(doc, (
        "Phase 3 validates each directional prediction — those classified as bullish "
        "or bearish with at least one resolvable ticker — against actual stock-price "
        "data fetched from Yahoo Finance. Two binary outcomes are computed for every "
        "validated prediction: whether the prediction was directionally correct "
        "(stock moved in the predicted direction), and whether it beat the passive "
        "market benchmark (the named stock out- or underperformed SPY over the same "
        "window)."
    ))

    add_heading2(doc, "5.1 Validation Methodology")
    add_body(doc, (
        "For each qualifying prediction the following steps are applied. First, the "
        "entry date is set to the first trading day on or after the tweet's "
        "created_date, ensuring we do not assume a trade executed before the tweet "
        "was public. Second, the target date is calculated as the entry date plus "
        "the horizon window in trading days (see §5.2). Third, look-ahead safety is "
        "enforced: predictions whose target date has not yet passed as of the run "
        "date are excluded rather than filled with partial data. Fourth, the "
        "adjusted-close prices of the stock and SPY at the entry and target dates "
        "are retrieved; 'adjusted close' accounts for stock splits and dividend "
        "distributions."
    ))
    add_body(doc, (
        "A prediction is judged correct if the stock's return over the window "
        "matches the predicted direction: positive return for bullish calls, "
        "negative return for bearish calls. A flat return (exactly zero) is "
        "treated as incorrect. A prediction beats the market if its excess return "
        "(stock return minus SPY return over the same window) is signed in the "
        "predicted direction: positive excess for bullish, negative excess for "
        "bearish. These two outcomes are independent — a bullish call on a stock "
        "that rises 2% when SPY rises 5% is directionally correct but does not "
        "beat the market."
    ))
    add_body(doc, (
        "Only the first ticker in tickers_mentioned is used for validation. "
        "This first-ticker rule is a deliberate simplification: when a tweet names "
        "multiple stocks it is not clear which one the prediction concerns most "
        "strongly. Using the first-mentioned ticker is documented and applied "
        "consistently; sensitivity to this choice is discussed in §5.4."
    ))

    add_heading2(doc, "5.2 Prediction Horizon Windows")
    add_body(doc, (
        "The time_horizon label assigned in Phase 2 maps to a fixed number of "
        "trading days:"
    ))
    add_body(doc, (
        "short_term → 21 trading days (approximately one calendar month). "
        "medium_term → 63 trading days (approximately one calendar quarter). "
        "long_term → 126 trading days (approximately six calendar months). "
        "unknown → 21 trading days (same as short_term)."
    ))
    add_body(doc, (
        "Tweets assigned horizon = 'unknown' by the classifier are validated using "
        "the 21-day window and flagged with horizon_was_unknown = True in the output "
        "dataset. This allows downstream analyses to exclude or stratify these rows "
        "if the arbitrary horizon assignment is a concern. The 21-day default is "
        "chosen because short_term is the modal horizon in the corpus (47.9% of "
        "classified tweets), and financial influencer content is disproportionately "
        "oriented toward short-term calls."
    ))

    add_heading2(doc, "5.3 Price Data and Caching")
    add_body(doc, (
        "Price histories are fetched from Yahoo Finance via the yfinance library "
        "(adjusted close, daily frequency, from 2020-01-01 to the run date). Each "
        "ticker's full history is downloaded once and cached to "
        "cache/prices/{TICKER}.csv. Subsequent runs read entirely from disk, making "
        "re-runs instantaneous and reproducible without internet access. Tickers "
        "that return an empty history (delisted symbols with no data, crypto tickers "
        "without the exchange suffix, etc.) receive a .notfound sentinel file so "
        "they are never retried. SPY is fetched and cached using the same mechanism."
    ))

    add_heading2(doc, "5.4 Validation Funnel")
    add_body(doc, (
        "Applying the validation pipeline to all 18,071 tweets yields the following "
        "funnel. Of 11,798 directional predictions (has_ticker AND sentiment in "
        "{bullish, bearish}), 1,102 were excluded because the first ticker returned "
        "an empty price history from Yahoo Finance — delisted stocks, crypto tokens "
        "without an exchange suffix (e.g. DOGE, SHIB, SOL), indices not available "
        "as Yahoo Finance tickers (e.g. SPX, VIX, DXY), and non-US symbols "
        "without US ADR listings (e.g. ADYEN, LVMH, TSMC). A further 4 predictions "
        "had a target date that had not yet passed as of the run date and 2 had "
        "no matching entry date in the price series. This leaves 10,690 validated "
        "predictions (90.6% of the directional set)."
    ))
    add_body(doc, (
        "Of the 10,690 validated predictions, 1,836 (17.2%) had time_horizon = "
        "'unknown' and were assigned the 21-day window by default; these are flagged "
        "via horizon_was_unknown = True in tweets_validated.csv and can be excluded "
        "in downstream analyses if sensitivity to the default assignment is a "
        "concern."
    ))

    # Funnel table
    add_body(doc, "Table 5.1. Data validation funnel from raw corpus to analysable predictions.")
    add_table(
        doc,
        headers=["Stage", "Count", "% of prior stage", "Notes"],
        rows=[
            ["All tweets in corpus",                  "18,071", "—",    ""],
            ["Has at least one ticker",                "14,039", "77.7%", "Regex + LLM extraction"],
            ["Directional (bullish or bearish)",       "11,798", "84.0%", "Neutral excluded"],
            ["Price data found (validated)",           "10,690", "90.6%", "Excl. 1,102 delisted/crypto; 6 other"],
        ],
        col_widths_cm=[5.5, 1.8, 2.8, 5.0],
        font_size_pt=9,
    )

    add_heading2(doc, "5.5 Prediction Accuracy")
    add_body(doc, (
        "Across all 10,690 validated predictions, 4,813 were directionally correct "
        "(45.0%) and 4,744 beat the SPY benchmark (44.4%). Both figures are "
        "materially below the 50% threshold expected from a random coin flip, "
        "suggesting that FinTwit predictions on average destroy value relative to "
        "passive buy-and-hold. A formal statistical test of significance is "
        "conducted in Phase 4 (RQ1)."
    ))
    add_body(doc, (
        "By sentiment: bullish calls (n = 8,083) were correct 40.5% of the time "
        "and beat the market 40.5% of the time. Bearish calls (n = 2,607) were "
        "correct 59.1% of the time and beat the market 56.3% of the time. The "
        "higher accuracy of bearish calls is consistent with the time period covered "
        "(2021-2025), which included the 2022 bear market during which many bearish "
        "predictions on individual stocks were vindicated. The bull-dominant sentiment "
        "distribution combined with sub-50% bullish accuracy largely explains the "
        "below-chance aggregate rate."
    ))
    add_body(doc, (
        "By prediction horizon: short_term (n = 5,774) correct 41.6%; "
        "medium_term (n = 1,869) correct 45.3%; long_term (n = 1,211) correct "
        "62.3%; unknown/defaulted-to-21d (n = 1,836) correct 44.2%. The long_term "
        "outperformance reflects the strong upward trend in US equities over the "
        "five-year study window: a bullish long-term call on most large-cap US "
        "stocks made at any point in 2021-2025 was more likely to be correct purely "
        "due to the overall market direction. This effect will be quantified and "
        "controlled for in the Phase 5 per-account analysis."
    ))


def write_phase4(doc, rq1: dict) -> None:
    add_heading1(doc, "6. RQ1 — Accuracy and Market Benchmark")
    add_body(doc, (
        "This section answers RQ1: how accurate are FinTwit predictions "
        "directionally, and do they beat a passive market benchmark? "
        "We assess two binary outcomes across all 10,690 validated predictions: "
        "prediction_correct (did the stock move in the predicted direction?) and "
        "beats_market (did the stock out- or underperform SPY over the same "
        "horizon?). We report 95% Wilson score confidence intervals throughout and "
        "test significance against a 50% null (a coin flip) using a two-tailed "
        "proportion z-test."
    ))

    add_heading2(doc, "6.1 Overall Accuracy")
    add_body(doc, (
        "Across all 10,690 validated predictions, 4,813 were directionally correct "
        "(45.0%; 95% CI [44.1%, 46.0%]). Against the market benchmark, 4,744 "
        "predictions beat SPY (44.4%; 95% CI [43.4%, 45.3%]). Both figures are "
        "statistically and practically significantly below the 50% benchmark expected "
        "from an uninformative coin flip."
    ))
    add_body(doc, (
        "Directional correct: z = -10.29, p = 7.75e-25. "
        "Beats market: z = -11.63, p = 3.05e-31. "
        "At a sample size of 10,690, these z-statistics are not merely crossing an "
        "arbitrary significance threshold — they reflect a systematic directional "
        "bias toward overconfident predictions. The aggregate portfolio of FinTwit "
        "calls actively destroys value relative to a passive buy-and-hold strategy "
        "in both the named stocks and the broad market."
    ))
    add_body(doc, (
        "Mean returns over the validated prediction windows: mean stock return "
        "-1.85% per prediction, mean SPY return -1.60%, mean excess return -0.24%. "
        "The negative mean stock and SPY returns reflect the 2022 market drawdown "
        "and the disproportionate short-term prediction horizon in the corpus (21 "
        "trading days, covered extensively during a high-volatility period). The "
        "negative mean excess return confirms the below-benchmark result on a "
        "magnitude basis, not just a directional one."
    ))

    add_heading2(doc, "6.2 Why beats_market Is the Cleaner Skill Measure")
    add_body(doc, (
        "Raw directional accuracy (prediction_correct) conflates genuine analytical "
        "skill with the unconditional tendency of equity markets to rise over time. "
        "A bullish call on any diversified portfolio of stocks held for 126 trading "
        "days (long_term horizon) will be correct simply because markets exhibit a "
        "positive long-run drift. The beats_market metric removes this base-rate "
        "advantage: it asks whether the prediction captured incremental information "
        "beyond what a passive SPY position would have produced. This distinction "
        "matters most for long-term predictions and is examined in §6.3."
    ))

    add_heading2(doc, "6.3 Accuracy by Segment")
    add_body(doc, (
        "Table 6.1 presents accuracy and beats_market rates with 95% Wilson CIs "
        "for each sub-group. Figure 6.1 shows the same results as a grouped bar "
        "chart with error bars. All sub-group estimates are below or at 50% except "
        "long_term prediction_correct (62.3%), which is explained below."
    ))
    add_body(doc, (
        "By sentiment — Bearish calls (n = 2,607) were correct 59.1% [57.2%, 60.9%] "
        "of the time and beat the market 56.3% [54.4%, 58.2%] of the time. Bullish "
        "calls (n = 8,083) were correct only 40.5% [39.4%, 41.6%] and beat the "
        "market 40.5% [39.5%, 41.6%]. The strong bearish outperformance on both "
        "metrics reflects the composition of the study period: the 2022 bear market "
        "vindicated many bearish predictions. Because bearish calls are a minority "
        "(22.1% of directional predictions), the overall aggregate rate is dragged "
        "down by the high-volume bullish calls, which performed substantially below "
        "chance. This is consistent with the promotional, hype-driven character of "
        "FinTwit content."
    ))
    add_body(doc, (
        "By prediction horizon — short_term: 41.6% correct, 41.1% beats "
        "(n = 5,774); medium_term: 45.3% correct, 46.7% beats (n = 1,869); "
        "long_term: 62.3% correct, 55.0% beats (n = 1,211); "
        "unknown (defaulted to 21d): 44.2% correct, 45.3% beats (n = 1,836). "
        "The long_term raw accuracy (62.3%) initially appears impressive but is "
        "substantially explained by the positive market drift over the 2021-2025 "
        "study window. The beats_market rate for long_term predictions drops to "
        "55.0% — above 50%, but much closer to chance, and not dramatically "
        "different from the medium_term beats_market rate of 46.7%. Short-term "
        "predictions — the plurality of the corpus — perform worst on both metrics, "
        "consistent with short-term price moves being nearly impossible to predict."
    ))
    add_body(doc, (
        "By trade type — analysis: 48.7% correct, 47.2% beats (n = 5,333); "
        "news: 47.0% correct, 46.7% beats (n = 758); "
        "general_discussion: 44.5% correct, 44.5% beats (n = 1,691); "
        "trade_suggestion: 38.0% correct, 38.5% beats (n = 2,908). "
        "The most striking finding here is that explicit trade recommendations — "
        "the subset of tweets that make a specific, actionable buy or sell call — "
        "are the worst performers by a considerable margin, 12 percentage points "
        "below the analysis category on both metrics. Analysis-type tweets (chart "
        "readings, reasoned opinion pieces) perform closest to 50% but still below "
        "it. The ordering is: analysis > news > general_discussion > "
        "trade_suggestion, exactly the reverse of what skill would predict."
    ))

    # Table 6.1
    seg_rows = _segment_table_rows(rq1)
    add_body(doc, "Table 6.1. Accuracy and beats_market rate by segment (95% Wilson CIs). "
                  "* p<0.05, ** p<0.01, *** p<0.001 vs 50% null (not shown here; see §6.1).")
    add_table(
        doc,
        headers=["Group", "n", "Correct %", "95% CI (correct)",
                 "Beats %", "95% CI (beats)"],
        rows=seg_rows,
        col_widths_cm=[4.2, 1.4, 1.8, 3.2, 1.8, 3.2],
        font_size_pt=8,
    )

    add_figure(doc, FIG_ACCURACY_SEG,
               "Figure 6.1. Prediction accuracy and beats_market rate by segment "
               "(sentiment, time horizon, trade type). "
               "Error bars are 95% Wilson confidence intervals. "
               "Dashed line marks 50% (no-skill baseline).")

    add_heading2(doc, "6.4 Sensitivity: Unknown-Horizon Rows")
    add_body(doc, (
        "Of the 10,690 validated predictions, 1,836 (17.2%) had time_horizon = "
        "'unknown' and were assigned the 21-day window by default (flagged as "
        "horizon_was_unknown). Excluding these rows yields n = 8,854 predictions. "
        "Results are virtually unchanged: directional correct 45.2% [44.2%, 46.2%] "
        "(p = 1.66e-19); beats market 44.2% [43.2%, 45.2%] (p = 8.75e-28). "
        "The headline finding is robust to the choice of a 21-day default for "
        "unspecified horizons."
    ))

    add_heading2(doc, "6.5 Accuracy by Calendar Year")
    add_figure(doc, FIG_ACCURACY_YEAR,
               "Figure 6.2. Directional accuracy and beats_market rate by calendar year "
               "(2021–2025), with 95% Wilson confidence intervals. "
               "Dashed line marks 50%.")
    add_body(doc, (
        "Figure 6.2 shows that directional accuracy and beats_market rates varied "
        "substantially by calendar year, tracking broad market conditions rather than "
        "any stable underlying skill: the 2022 bear market depressed both metrics for "
        "the corpus's bullish-dominant predictions, while subsequent recovery years "
        "partially restored them. This regime-driven variation — not persistent "
        "analyst ability — is the primary explanation for year-on-year performance "
        "swings, a confound that the per-account shrinkage model in §7 must isolate."
    ))

    add_heading2(doc, "6.6 Interpretation")
    add_body(doc, (
        "The answer to RQ1 is unambiguous: FinTwit predictions do not beat chance "
        "and do not beat the market. Both headline metrics — directional accuracy "
        "and beats_market rate — are significantly below 50%, with p-values on the "
        "order of 10^-25 to 10^-31. The margin of underperformance (-5 to -6 "
        "percentage points relative to a coin flip) is statistically robust to "
        "segmentation, sensitivity checks, and choice of outcome metric."
    ))
    add_body(doc, (
        "The sub-group analysis reveals two notable patterns that will be explored "
        "further in RQ2 (per-account skill) and RQ3 (signal value). First, bearish "
        "calls outperform bullish calls by ~18 percentage points, likely driven by "
        "the 2022 bear market period rather than inherent skill — a hypothesis to "
        "be tested in RQ2 by controlling for market regime. Second, the most "
        "explicit, actionable call type (trade_suggestion) is the worst-performing "
        "category, suggesting that increased conviction in the tweet's language does "
        "not map to increased accuracy — a relationship formally examined in RQ3 "
        "via the LLM's own confidence score."
    ))

    add_figure(doc, FIG_EXCESS_RET,
               "Figure 6.3. Distribution of per-prediction excess return (stock return minus "
               "SPY return, clipped at ±0.5 for display). "
               "The mean excess return of –0.24% confirms below-benchmark performance "
               "on a magnitude basis.")


def write_phase5(doc, rq2: dict) -> None:
    add_heading1(doc, "7. RQ2 — Skill vs Luck")
    add_body(doc, (
        "RQ2 asks whether any individual account's track record is statistically "
        "distinguishable from chance after accounting for the multiple-comparisons "
        "problem inherent in evaluating 47 accounts simultaneously. We use the primary "
        "skill metric beats_market (§6.2) and apply both Benjamini-Hochberg FDR and "
        "Bonferroni correction. An empirical-Bayes shrinkage model then quantifies "
        "how much of each account's apparent performance is real signal vs noise."
    ))

    add_heading2(doc, "7.1 Account Eligibility and Method")
    add_body(doc, (
        "Of the 51 pseudonymised accounts in the corpus, 50 contributed at least one "
        "validated prediction. We set a minimum of 30 validated predictions to ensure "
        "that per-account binomial tests have at least nominal power; below this "
        "threshold, a 95% confidence interval on a proportion is too wide to be "
        "informative. Three accounts fall below this threshold (n = 10, 10, 20) and "
        "are excluded from the significance analysis, leaving 47 qualifying accounts "
        "with n ranging from 35 to 566 validated predictions."
    ))
    add_body(doc, (
        "For each qualifying account we run a two-sided exact binomial test of "
        "H0: p_beats = 0.50 (no skill). The two-sided framing is deliberate: an "
        "account that beats the market far less than 50% of the time is also "
        "informative — consistently wrong predictions in a predictable direction "
        "could, in principle, be exploited as a contrarian signal."
    ))

    add_heading2(doc, "7.2 Multiple-Testing Correction — The Headline")
    add_body(doc, (
        "Before any correction, 24 of 47 accounts are nominally significant at "
        "p < 0.05: 9 above 50% (apparently skilled) and 15 below 50% (apparently "
        "anti-skilled). With 47 tests at alpha = 0.05, we expect roughly 47 × 0.05 "
        "= 2.4 false positives by chance alone, so the raw count of 24 is far more "
        "than the false-discovery rate can explain in isolation."
    ))
    add_body(doc, (
        "After Benjamini-Hochberg FDR correction (alpha = 0.05): 23 accounts survive, "
        "9 above 50% and 14 below 50%. BH-FDR controls the expected proportion of "
        "false discoveries among the rejected hypotheses and is the recommended "
        "procedure for exploratory testing of this kind."
    ))
    add_body(doc, (
        "After Bonferroni correction (alpha = 0.05 / 47 = 0.00106): 11 accounts "
        "survive, 1 above 50% and 10 below 50%. Bonferroni controls the family-wise "
        "error rate (probability of any false positive) and is substantially stricter. "
        "The key contrast is that under the lenient FDR criterion 9 accounts appear "
        "genuinely skilled at beating the market, but this collapses to 1 account "
        "under the strict family-wise control."
    ))
    add_body(doc, (
        "The single account that survives Bonferroni correction with beats_market > 50% "
        "is account_19 (n = 83, raw rate 69.9%, shrunk estimate 66.1%, p = 0.0004, "
        "p_Bonferroni = 0.018). Ten accounts survive Bonferroni with beats_market "
        "significantly below 50%, the worst being account_01 (n = 456, raw rate 4.2%, "
        "p << 0.001). The consistently anti-skilled accounts are as statistically "
        "remarkable as any candidate 'skill' — they represent systematic, predictable "
        "wrongness rather than random noise."
    ))

    add_figure(doc, FIG_CATERPILLAR,
               "Figure 7.1. Per-account beats_market rate for all 47 qualifying accounts "
               "(sorted by shrunk posterior mean). Filled circles = shrunk (empirical-Bayes) "
               "estimate; hollow circles = raw rate; bars = 95% credible interval. "
               "Red = CI entirely above 50%; orange dotted line = prior mean.")
    add_figure(doc, FIG_MULTTEST,
               "Figure 7.2. Number of accounts achieving statistical significance before and "
               "after multiple-testing correction, split by direction (above vs below 50%). "
               "Bonferroni collapses the ‘skilled’ count from 9 to 1.")

    add_heading2(doc, "7.3 Empirical-Bayes Shrinkage")
    add_body(doc, (
        "We fit a Beta prior from the pooled per-account rates using the method of "
        "moments, yielding Beta(7.76, 8.76) with prior mean 47.0% — close to the "
        "overall corpus beats_market rate of 44.4%, as expected. The posterior for "
        "each account is Beta(alpha_prior + k_beats, beta_prior + n - k_beats), giving "
        "a shrunken estimate of each account's true beats_market rate."
    ))
    add_body(doc, (
        "The shrinkage effect is strong for accounts with small n and weak for accounts "
        "with large n, as the theory requires. Account_37's raw rate (74.3%, n = 35) "
        "shrinks to 65.5% after pooling with the prior; account_01's raw rate (4.2%, "
        "n = 456) barely moves in the shrinkage — 5.7% — because the large sample "
        "overpowers the prior. Table 7.1 lists full results for all 47 qualifying "
        "accounts sorted by shrunk estimate."
    ))
    add_body(doc, (
        "Nine accounts have a 95% credible interval for beats_market that lies entirely "
        "above 50%: account_19, account_37, account_45, account_03, account_09, "
        "account_13, account_39, account_29, and account_06. These accounts appear "
        "genuinely skilled under the shrinkage model — their posterior distributions "
        "are inconsistent with the null of 50% even after pooling toward the prior. "
        "Notably, these are the same 9 that survive BH-FDR correction."
    ))

    # Table 7.1
    acct_rows = _account_table_rows(rq2)
    add_body(doc, "Table 7.1. Per-account beats_market results for all 47 qualifying accounts "
                  "(sorted by shrunk estimate, descending). "
                  "↑ = 95% credible interval entirely above 50%. "
                  "* p<0.05, ** p<0.01, *** p<0.001 (raw p-value).")
    add_table(
        doc,
        headers=["Account", "n", "Raw beats%", "Shrunk%",
                 "95% Cred. interval", "p (raw)", "p (BH)", "p (Bonf.)"],
        rows=acct_rows,
        col_widths_cm=[2.3, 1.0, 1.7, 1.7, 3.5, 1.5, 1.5, 1.5],
        font_size_pt=8,
    )

    add_heading2(doc, "7.4 Interpretation")
    add_body(doc, (
        "The answer to RQ2 is nuanced rather than a simple 'no skill survives'. "
        "Under BH-FDR correction — appropriate for exploratory research — 9 of 47 "
        "accounts (19%) show beats_market rates credibly above 50%. Under the "
        "stricter Bonferroni criterion, only 1 account survives with a rate above 50%. "
        "The full picture is therefore: a heterogeneous population where the majority "
        "of accounts perform near or below the market benchmark, a clear tail of "
        "accounts that are systematically wrong (likely due to consistent bullish "
        "bias during the 2022 bear market), and a smaller tail of accounts that "
        "appear to add value — though whether this reflects genuine analytical skill "
        "or style exposure (e.g., persistent bearish positioning during a bear market "
        "period) cannot be determined from the current analysis."
    ))
    add_body(doc, (
        "Phase 6 (RQ3) examines whether observable tweet-level signals — trade type, "
        "confidence, engagement — predict which predictions come true, irrespective "
        "of which account made them. This provides a complementary and more "
        "generalisable test of whether any systematic predictive structure exists "
        "in the corpus beyond account identity."
    ))


def write_phase6(doc, calib: pd.DataFrame) -> None:
    add_heading1(doc, "8. RQ3 — Signal Value and Calibration")
    add_body(doc, (
        "RQ3 asks whether tweet-level signals — observable at or shortly after the "
        "time of posting — carry any information about whether a prediction will come "
        "true. We assess this in two complementary ways: (a) a formal calibration "
        "analysis of the LLM's self-reported confidence score, and (b) a predictive "
        "model of prediction_correct using all available tweet-time features."
    ))

    add_heading2(doc, "8.1 Confidence Calibration")
    add_body(doc, (
        "The LLM classifier (Phase 2) assigned each tweet a confidence score from "
        "0 to 100 reflecting how assertively the tweet was written. In a well-calibrated "
        "system, predictions at confidence 80 should be correct approximately 80% of the "
        "time. We assess this via a reliability curve (Figure 8.1) and a Pearson "
        "correlation."
    ))
    add_body(doc, (
        "Reliability curve — Across the 10,690 validated predictions, the "
        "distribution of confidence scores is extremely narrow (mean 76.3, "
        "standard deviation 6.2, range 20-100). The vast majority of predictions "
        "cluster in the 70-80 and 80-90 bins, yielding only six non-empty bins. "
        "Actual prediction_correct rates within each bin are: 55.0% [60-70), "
        "44.7% [70-80), 44.7% [80-90), 52.0% [90-100). There is no systematic "
        "upward trend — a more assertively-written prediction is not more likely "
        "to be correct."
    ))
    add_body(doc, (
        "Pearson correlation — The point-biserial correlation between confidence "
        "and prediction_correct is r = -0.0005, p = 0.9624 (95% CI [-0.019, +0.019]). "
        "The correlation with beats_market is r = +0.0084, p = 0.3870 (n = 10,690). "
        "Both correlations are indistinguishable from zero. The LLM's confidence score "
        "is completely uncalibrated: it reflects assertive linguistic register rather "
        "than predictive validity."
    ))

    add_figure(doc, FIG_CALIB_CURVE,
               "Figure 8.1. Reliability curve: LLM confidence score (x-axis) vs actual "
               "prediction_correct rate (y-axis). Each point is a 10-point confidence bin; "
               "error bars are 95% Wilson CIs. The flat line confirms near-zero calibration "
               "(r = −0.0005, p = 0.96).")

    # Table 8.1
    calib_rows = _calib_table_rows(calib)
    add_body(doc, "Table 8.1. Confidence calibration table: per-bin actual accuracy "
                  "with 95% Wilson confidence intervals.")
    add_table(
        doc,
        headers=["Confidence bin", "n", "Actual accuracy", "95% CI"],
        rows=calib_rows,
        col_widths_cm=[3.0, 1.8, 2.8, 3.5],
        font_size_pt=9,
    )

    add_figure(doc, FIG_CONF_HIST,
               "Figure 8.4. Distribution of LLM confidence scores across the 10,690 "
               "validated predictions. The narrow, high-skew distribution (mean 76.3, "
               "SD 6.2, range 20–100) leaves little variation for calibration analysis.")

    add_heading2(doc, "8.2 Predictive Model")
    add_body(doc, (
        "We train two classification models to predict prediction_correct using tweet-time "
        "features only. Features are: sentiment (is_bullish dummy, reference = bearish); "
        "time horizon (medium_term, long_term, unknown dummies, reference = short_term); "
        "trade type (trade_suggestion, general_discussion, news dummies, reference = "
        "analysis); confidence (continuous); and log1p-transformed engagement counts "
        "(log_likes, log_retweets, log_replies_count, log_views, log_author_followers). "
        "No account-level accuracy history is included — that would constitute data "
        "leakage from RQ2. Model evaluation uses 5-fold stratified cross-validation; "
        "the primary metric is AUC-ROC (baseline 0.500 for a random classifier; "
        "majority-class accuracy baseline is 55.0%)."
    ))
    add_body(doc, (
        "Logistic regression (L2, StandardScaler): mean CV AUC = 0.654 "
        "(delta vs baseline = +0.154). "
        "Gradient boosting (100 trees, max_depth=3, learning_rate=0.05): "
        "mean CV AUC = 0.700 (delta vs baseline = +0.200). "
        "Both models meaningfully outperform the random baseline, indicating that the "
        "observable tweet-level signals do carry real predictive information."
    ))

    add_figure(doc, FIG_ROC,
               "Figure 8.3. ROC curves for logistic regression and gradient boosting, "
               "computed from 5-fold stratified cross-validation out-of-fold probability "
               "predictions. Both models substantially outperform the random baseline "
               "(diagonal dashed line).")

    add_heading2(doc, "8.3 Feature Importance")
    add_body(doc, (
        "Logistic regression standardised coefficients (sorted by magnitude) — "
        "The three dominant features are: log_likes (+0.424), is_bullish (-0.412), "
        "and horizon_long_term (+0.266). Negative is_bullish coefficient confirms "
        "that bearish predictions are more accurate (consistent with RQ1: bearish "
        "accuracy 59.1% vs bullish 40.5%). The positive coefficient for horizon_long_term "
        "confirms the long-term drift effect (62.3% raw accuracy, largely explained by "
        "the market's positive trend over 2021-2025). Higher engagement (log_likes) "
        "is positively associated with accuracy — predictions that attracted more likes "
        "tended to come true, though this metric accumulates after posting and is not "
        "available at prediction time for real-time use. Confidence appears near zero "
        "(coef = +0.048, OR = 1.05) — consistent with the calibration finding."
    ))
    add_body(doc, (
        "Gradient boosting feature importances (impurity decrease) — "
        "log_author_followers dominates (0.429), followed by is_bullish (0.235) and "
        "horizon_long_term (0.190). The prominence of log_author_followers in the GB "
        "model suggests that account size is a strong proxy for prediction style: "
        "larger accounts in this corpus tend toward bullish, short-term trade suggestions "
        "— the patterns associated with lower accuracy in RQ1 — while smaller accounts "
        "show more heterogeneous styles. Confidence ranks ninth (0.004), confirming "
        "that the LLM's self-assessed certainty carries almost no predictive weight "
        "beyond the categorical features already captured."
    ))

    add_figure(doc, FIG_FEAT_IMPORT,
               "Figure 8.2. Feature importance: logistic regression standardised coefficients "
               "(left) and gradient boosting impurity-decrease importances (right). "
               "Positive LR coefficients increase the probability of prediction_correct. "
               "Confidence ranks near the bottom in both models.")

    add_heading2(doc, "8.4 Interpretation")
    add_body(doc, (
        "The answer to RQ3 is two-part. First, the LLM's confidence score is entirely "
        "uncalibrated (r ≈ 0). High-confidence tweets are not more likely to come true. "
        "This is expected: the confidence score captures assertive language (which is "
        "endemic in FinTwit), not analytical quality. Second, tweet-level signals do "
        "carry predictive information: both models achieve AUC ≈ 0.65–0.70, substantially "
        "above the 0.50 baseline."
    ))
    add_body(doc, (
        "Crucially, however, the informative features are almost entirely the categorical "
        "signals already characterised in RQ1 — sentiment direction and prediction horizon "
        "— plus a post-hoc engagement proxy (log_likes) that is not available in real "
        "time. The model is largely rediscovering the RQ1 base rates: bearish calls are "
        "more accurate (2022 bear market), long-term calls appear more accurate (market "
        "drift), and explicit trade suggestions are less accurate. It is not uncovering "
        "new, orthogonal signal. A practically useful real-time screening model would "
        "need features available before the prediction outcome is known — and of those, "
        "only sentiment and horizon provide meaningful lift."
    ))
    add_body(doc, (
        "Taken together with RQ1 and RQ2, the full picture is: FinTwit predictions "
        "systematically underperform a coin flip in aggregate; the heterogeneity across "
        "accounts is real but mostly reflects market-regime exposure rather than "
        "persistent skill; and no tweet-level signal (including LLM confidence) is "
        "sufficiently calibrated or strong to identify which individual predictions will "
        "come true. The most actionable implication for a retail investor following "
        "FinTwit accounts is the contrarian one: high-volume, high-conviction bullish "
        "trade suggestions from large accounts have been the single worst-performing "
        "category in this corpus."
    ))


def _regime_skilled_rows(regime: dict) -> list[list[str]]:
    """Build Table 9.1 rows from the skilled-account regime table."""
    import math
    tbl = regime["skilled"]
    rows = []
    for acct, r in tbl.iterrows():
        p1_rate_s = (f"{r['p1_beats_rate']*100:.1f}%"
                     if not math.isnan(r["p1_beats_rate"]) else "N/A")
        p1_ci_s   = (f"[{r['p1_ci_low']*100:.1f}%, {r['p1_ci_high']*100:.1f}%]"
                     if not math.isnan(r["p1_ci_low"]) else "N/A")
        p2_rate_s = f"{r['p2_beats_rate']*100:.1f}%"
        p2_ci_s   = f"[{r['p2_ci_low']*100:.1f}%, {r['p2_ci_high']*100:.1f}%]"
        both_s    = "Yes†" if r["above_50_both"] else ("—" if r["p1_n"] == 0 else "No")
        bonf_flag = " *" if acct == "account_19" else ""
        rows.append([acct + bonf_flag,
                     str(int(r["p1_n"])), p1_rate_s, p1_ci_s,
                     str(int(r["p2_n"])), p2_rate_s, p2_ci_s,
                     both_s])
    return rows


def write_phase7(doc, regime: dict) -> None:
    """Section 9 — Robustness: Skill vs Style Across Regimes."""
    split = regime["split"]
    pers  = regime["persistence"]
    p1    = split["p1"]
    p2    = split["p2"]

    add_heading1(doc, "9. Robustness — Skill vs Style Across Regimes")
    add_body(doc, (
        "Sections 6 and 7 established that a subset of accounts outperform the "
        "50% beats_market threshold at the aggregate level (2021-2025 combined). "
        "A central question is whether this outperformance reflects durable analytical "
        "skill or merely style alignment with the 2023-2025 bull market that dominates "
        "the corpus. The design intended a split-half persistence test using the 2022 "
        "bear market as a held-out regime. The primary finding of this section is that "
        "the test is infeasible: the data are too heavily concentrated in the post-2022 "
        "period for a formal cross-regime comparison to be conducted. This concentration "
        "is itself a substantive limitation on the §7 skill claims."
    ))

    add_heading2(doc, "9.1 Method")
    add_body(doc, (
        "Validated predictions are split by entry date into two sub-periods. "
        "P1 (Jan 2021 – Dec 2022) encompasses the bull market of 2021 and the "
        "sharp 2022 bear market. P2 (Jan 2023 – Feb 2025) covers the post-bear "
        "recovery and the strong bull market of 2024-2025. The primary skill metric "
        "is beats_market, consistent with §7. The intended persistence test computes "
        "the Pearson and Spearman rank correlation between per-account P1 and P2 "
        "beats_market rates, restricted to accounts with at least 20 validated "
        "predictions in each sub-period individually."
    ))

    add_heading2(doc, "9.2 Sub-Period Data Overview")
    add_body(doc, (
        f"P1 contains {p1['n']:,} validated predictions ({100*p1['n']/(p1['n']+p2['n']):.1f}% "
        f"of the total); P2 contains {p2['n']:,} ({100*p2['n']/(p1['n']+p2['n']):.1f}%). "
        f"The corpus is overwhelmingly concentrated in P2: 97.9% of validated predictions "
        f"have entry dates in 2023 or later, reflecting when the FinTwit accounts in "
        f"this dataset were most active. The year-by-year breakdown is: 2021 n=9, "
        f"2022 n=51, 2023 n=307, 2024 n=3,747, 2025 n=6,575 (plus 1 in early 2026 "
        f"due to a long horizon window)."
    ))
    add_body(doc, (
        f"P1 beats_market rate: {p1['beats_rate']*100:.1f}% "
        f"(95% CI [{p1['beats_ci'][0]*100:.1f}%, {p1['beats_ci'][1]*100:.1f}%], n={p1['n']}). "
        f"P2 beats_market rate: {p2['beats_rate']*100:.1f}% "
        f"(95% CI [{p2['beats_ci'][0]*100:.1f}%, {p2['beats_ci'][1]*100:.1f}%], n={p2['n']:,}). "
        f"The P1 confidence interval is wide (n={p1['n']} only) and uninformative on its own. "
        f"The P2 rate is consistent with the overall 44.4% reported in §6.1."
    ))

    add_heading2(doc, "9.3 Persistence Test — Finding: Infeasible")
    pers_acc = pers["account_df"]
    if not pers_acc.empty:
        acct_list = "; ".join(
            f"{acct} (P1 n={int(row['p1_n'])}, P2 n={int(row['p2_n'])})"
            for acct, row in pers_acc.iterrows()
        )
        acct_summary = (
            f"Only {pers['n_qualifying']} of 51 accounts have at least 20 validated "
            f"predictions in both P1 and P2: {acct_list}."
        )
    else:
        acct_summary = "No accounts have at least 20 validated predictions in both sub-periods."
    add_body(doc, (
        f"The split-half persistence correlation cannot be computed from this corpus. "
        f"{acct_summary} "
        f"A minimum of 4 qualifying accounts is required for a Pearson or Spearman "
        f"correlation to carry any statistical power; with fewer than 4, the statistic "
        f"is degenerate and the p-value meaningless. The formal test is therefore not run."
    ))
    add_body(doc, (
        "The deeper finding is structural: 99.2% of validated predictions post-date "
        "2022, meaning the 2022 bear market is effectively absent from the corpus as a "
        "testing regime. Most accounts either began posting after 2022 or had their "
        "P1-era predictions excluded because the relevant tickers were unavailable in "
        "Yahoo Finance (delisted stocks, crypto, etc.). Critically, 8 of the 9 accounts "
        "identified as 'skilled' in §7 have zero validated predictions in P1. Their "
        "entire track records — and therefore their §7 skill designations — are drawn "
        "from P2 (2023-2025) alone: a predominantly bullish market period in which a "
        "systematically bullish account will automatically accumulate a positive record. "
        "It is not possible to distinguish genuine analytical skill from bullish style "
        "alignment on the basis of P2-only data."
    ))

    add_figure(doc, FIG_REGIME,
               "Figure 9.1. P1 (Jan 2021 – Dec 2022) vs P2 (Jan 2023 – Feb 2025) "
               "beats_market rate for the two accounts with ≥20 validated predictions "
               "in both sub-periods. The 45-degree line indicates perfect persistence; "
               "red dotted lines mark the 50% chance level. Note the severe data "
               "imbalance: 99.2% of validated predictions fall in P2.")

    add_heading2(doc, "9.4 Skilled Accounts Under Regime Scrutiny")
    add_body(doc, (
        "Table 9.1 presents the P1 and P2 beats_market rates for the nine accounts "
        "whose 95% credible intervals were entirely above 50% in §7. Account_19 "
        "(marked * as the sole Bonferroni survivor) is the only account with any "
        "P1 data. The remaining eight accounts have zero validated predictions in P1: "
        "their §7 track records are based entirely on P2 (2023-2025) performance."
    ))

    skilled_rows = _regime_skilled_rows(regime)
    add_body(doc,
             "Table 9.1. Beats_market rate and 95% Wilson CI in P1 and P2 for the nine "
             "accounts credibly above 50% in §7 (sorted by §7 ranking). "
             "* = sole Bonferroni survivor (§7). † = above 50% in both periods. "
             "— = no P1 data available.")
    add_table(
        doc,
        headers=["Account", "P1 n", "P1 beats%", "P1 95% CI",
                 "P2 n", "P2 beats%", "P2 95% CI", "Both>50%"],
        rows=skilled_rows,
        col_widths_cm=[2.4, 1.0, 1.8, 3.0, 1.0, 1.8, 3.0, 1.8],
        font_size_pt=8,
    )
    add_body(doc, (
        "Account_19 (the sole Bonferroni survivor from §7) is the only skilled-list "
        "account with any P1 data. Its P1 beats_market rate of 87.5% "
        "(95% CI [71.9%, 95.0%], n=32) is high, though the sample is small. "
        "In P2 (n=51), its rate falls to 58.8%, with a 95% CI of [45.2%, 71.2%] "
        "that spans 50%: the P2 point estimate is above chance but is not "
        "statistically distinguishable from it. The 'Both>50%' flag in the table "
        "is a point-estimate observation, not a significance test. Taken together, "
        "account_19's record provides the only cross-regime data available in this "
        "corpus, but it does not constitute evidence of durable, statistically robust "
        "skill: the stronger result is in P1 (a small sample) and the recent-period "
        "result is not significant. For the remaining eight accounts, no such check "
        "is possible at all."
    ))

    add_heading2(doc, "9.5 Verdict")
    add_body(doc, (
        "The regime analysis yields three findings. First, the split-half persistence "
        "test is infeasible: 99.2% of validated predictions fall in P2, leaving P1 "
        "too thin for a formal correlation. This is not a contingent failure — it "
        "reflects an irreversible feature of the data collection window. Second, "
        "8 of the 9 §7 'skilled' accounts have their entire track record in P2 "
        "(2023-2025). Their §7 skill designations cannot be separated from bullish "
        "style alignment in a predominantly up-trending market; regime-robustness "
        "cannot be assessed for them at all. Third, account_19 — the one account "
        "with cross-regime data — shows a high P1 point estimate (87.5%) but a P2 "
        "confidence interval [45.2%, 71.2%] that includes 50%, so its recent-period "
        "performance is not distinguishable from chance. The overall conclusion is "
        "that no account in this corpus demonstrates durable skill that survives a "
        "regime change: the data do not permit the test for 8 of 9 candidates, and "
        "the one testable account loses statistical significance in the later period."
    ))


def write_limitations(doc) -> None:
    """Section 10 — Limitations."""
    add_heading1(doc, "10. Limitations")
    add_body(doc, (
        "This analysis is rigorous within its design, but several structural "
        "limitations bound the conclusions."
    ))

    add_heading2(doc, "10.1 Survivorship and Selection Bias")
    add_body(doc, (
        "The dataset captures only accounts and tweets that were public and "
        "accessible at the time of the scrape. Deleted tweets, suspended accounts, "
        "and accounts that went private before the collection date are entirely "
        "absent. Accounts that had poor track records and subsequently deleted their "
        "content or deactivated would systematically inflate the apparent performance "
        "of the observed corpus — a classic survivorship-bias mechanism. The direction "
        "of this bias is toward over-estimating accuracy and over-counting skilled "
        "accounts. Given the corpus already shows below-50% aggregate accuracy, "
        "survivorship bias would make the true population performance even worse."
    ))

    add_heading2(doc, "10.2 First-Ticker Simplification")
    add_body(doc, (
        "When a tweet mentions multiple tickers, only the first is used for "
        "price validation. This is documented and applied consistently, but it "
        "ignores multi-stock predictions and potentially misattributes sentiment "
        "to the wrong ticker. Sensitivity to this choice is not formally tested; "
        "results could differ if the second or most-prominent ticker were used instead."
    ))

    add_heading2(doc, "10.3 LLM Classification Error")
    add_body(doc, (
        "The sentiment, time_horizon, trade_type, and confidence fields are "
        "gpt-4o-mini outputs, not ground truth. Classification errors are not "
        "propagated into the statistical analysis — the model's outputs are treated "
        "as if they were exact. A formal precision/recall evaluation was performed "
        "for ticker extraction only (§3.5); no equivalent evaluation was done for "
        "the four classification fields. Misclassification would add noise to the "
        "segmented analyses in §6.3, bias the calibration curves in §8.1, and "
        "contaminate feature importance estimates in §8.3."
    ))

    add_heading2(doc, "10.4 beats_market Practicality")
    add_body(doc, (
        "The beats_market metric computes excess return relative to SPY assuming "
        "execution at the adjusted close on the first trading day on or after the "
        "tweet date. In practice, a retail investor faces bid-ask spreads, brokerage "
        "commissions, and potential price impact. For bearish predictions, the "
        "metric implicitly assumes short-selling feasibility, which is not always "
        "possible (hard-to-borrow stocks, margin requirements). These implementation "
        "frictions would reduce or eliminate most of the apparent gains reported "
        "for any 'skilled' account, and make the already-poor aggregate performance "
        "even worse in a real-trading context."
    ))

    add_heading2(doc, "10.5 Fixed Horizon Windows and the Unknown Default")
    add_body(doc, (
        "The time-horizon mapping is a deliberate, documented simplification: "
        "short_term = 21 trading days, medium_term = 63, long_term = 126. These "
        "are reasonable conventions but do not precisely match every account's "
        "implied horizon. The 1,836 predictions (17.2%) assigned horizon = 'unknown' "
        "are validated at 21 days by default. Section 6.4 shows that excluding "
        "these rows does not materially change the headline results, but the "
        "default assignment could still bias within-horizon comparisons."
    ))

    add_heading2(doc, "10.6 Exclusion of Neutral Tweets")
    add_body(doc, (
        "Tweets classified as sentiment = 'neutral' (2,241 tweets, 16.0% of "
        "classified tweets) are excluded from all price-validation and accuracy "
        "analyses. Neutral tweets may carry useful information (e.g., tone shifts, "
        "sector commentary) that is not captured. The analysis therefore speaks "
        "only to explicitly directional calls."
    ))

    add_heading2(doc, "10.7 Period Specificity")
    add_body(doc, (
        "The study covers January 2021 to February 2025 — a window that includes "
        "the 2021 bull market, the 2022 bear market, and the 2023-2025 recovery. "
        "While this multi-year span adds breadth, the data is heavily concentrated "
        "in 2024-2025 (§9.2). Conclusions about the 2022 bear-market period are "
        "based on very few observations. Performance patterns that appear robust "
        "within this window may not generalise to other market regimes or time periods. "
        "The regime analysis in §9 explicitly addresses this concern and concludes "
        "that a cross-regime persistence test is not feasible with the available data."
    ))

    add_heading2(doc, "10.8 Data Concentration and the Limits of Skill Inference")
    add_body(doc, (
        "99.2% of validated predictions fall in P2 (Jan 2023 – Feb 2025). This "
        "concentration has a direct and underappreciated consequence for the skill "
        "claims in §7: no account identified as 'skilled' can be shown to perform "
        "above chance outside the 2023-2025 bull market, because no meaningful P1 "
        "data exists for them. The §7 analysis correctly controls for multiple "
        "comparisons within the available data, but it cannot control for the "
        "possibility that above-50% performance in 2023-2025 reflects regime exposure "
        "rather than analytical ability. Any inference about individual account skill "
        "must therefore be read with this constraint: the data are insufficient to "
        "distinguish skill from bullish style in a rising market."
    ))


def write_conclusion(doc) -> None:
    """Section 11 — Conclusion."""
    add_heading1(doc, "11. Conclusion")
    add_body(doc, (
        "This project conducted a large-scale, reproducible statistical audit of "
        "~18,000 stock predictions made by 51 financial-influencer accounts on X "
        "between January 2021 and February 2025. The analysis addressed three "
        "pre-specified research questions and a robustness check. The findings are "
        "internally consistent and their direction is unambiguous, even if the "
        "magnitude of heterogeneity varies by subgroup."
    ))

    add_heading2(doc, "11.1 RQ1 — Aggregate Underperformance")
    add_body(doc, (
        "Across all 10,690 validated directional predictions, 45.0% were "
        "directionally correct and 44.4% beat the SPY benchmark. Both figures are "
        "statistically and practically significantly below the 50% baseline expected "
        "from a random coin flip (z ≈ −10 to −12, p < 10^−24). This aggregate "
        "underperformance is robust to subgroup splitting, the exclusion of "
        "unknown-horizon rows, and the choice of metric (correct vs beats_market). "
        "FinTwit predictions, taken as a portfolio, actively destroy value relative "
        "to a passive buy-and-hold strategy."
    ))

    add_heading2(doc, "11.2 RQ2 — Heterogeneity, but No Durable Skill")
    add_body(doc, (
        "The aggregate result masks substantial heterogeneity across accounts. After "
        "Benjamini-Hochberg FDR correction, nine accounts show beats_market rates "
        "credibly above 50%; after the stricter Bonferroni correction, only one "
        "(account_19) survives. The empirical-Bayes shrinkage model confirms that "
        "some of this dispersion reflects real differences rather than sampling noise. "
        "The regime analysis in §9, however, forces a harder conclusion: no account "
        "in this corpus demonstrates durable, statistically robust skill that survives "
        "a market-regime change. For 8 of the 9 'skilled' accounts, the regime test "
        "cannot even be attempted — their entire track records fall within the "
        "2023-2025 bull market, making it impossible to separate skill from bullish "
        "style alignment. For the one testable account (account_19), the P2 "
        "confidence interval [45.2%, 71.2%] includes 50%, meaning its recent-period "
        "performance is not statistically distinguishable from chance. The bottom "
        "line on RQ2 is therefore negative: apparent outperformance exists within "
        "the sample, but there is no evidence that it is regime-robust or would "
        "persist out of sample."
    ))

    add_heading2(doc, "11.3 RQ3 — Uncalibrated Confidence, Modest Signal")
    add_body(doc, (
        "The LLM's self-reported confidence score is completely uncalibrated "
        "(r = −0.0005 vs prediction_correct, p = 0.96). High-confidence tweets are "
        "no more likely to come true than low-confidence ones. This is expected: the "
        "score captures assertive linguistic register, not analytical quality, and "
        "the distribution is so narrow (mean 76.3, SD 6.2) that it lacks the "
        "variation needed for calibration. Tweet-level features do carry some "
        "predictive signal — logistic regression achieves AUC ≈ 0.65 and gradient "
        "boosting AUC ≈ 0.70 — but the informative features are almost entirely "
        "the categorical signals already characterised in RQ1 (sentiment direction, "
        "horizon, trade type). The model is rediscovering RQ1 base rates, not "
        "uncovering novel structure."
    ))

    add_heading2(doc, "11.4 Contrarian Takeaway")
    add_body(doc, (
        "The most practically significant finding for a retail investor who follows "
        "FinTwit accounts is the contrarian one: high-volume, high-conviction "
        "bullish trade suggestions from large accounts represent the single "
        "worst-performing category in this corpus (38.0% correct, 38.5% beats "
        "market). The features that make a tweet maximally persuasive — explicit "
        "trade call, bullish framing, posted by a large account, high LLM confidence "
        "score — are negatively associated with accuracy. Following these signals "
        "would have underperformed a random coin flip by roughly 12 percentage points."
    ))

    add_heading2(doc, "11.5 Final Note")
    add_body(doc, (
        "This analysis is a descriptive and inferential audit of a specific corpus "
        "over a specific period. It does not constitute financial advice, and the "
        "findings may not generalise to other platforms, other time periods, or the "
        "broader population of financial social-media accounts. The code, data, and "
        "statistical choices are fully documented and reproducible; the conclusions "
        "are as strong as the data allow — which is to say: strongly negative on "
        "aggregate skill, heterogeneous at the account level, and honest about the "
        "limits of what the available data can establish."
    ))


# ── main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    if not INPUT_CSV.exists():
        raise SystemExit(f"ERROR: {INPUT_CSV} not found.")

    print(f"Loading {INPUT_CSV} ...")
    df = pd.read_csv(INPUT_CSV, low_memory=False)

    # ── generate all figures ──────────────────────────────────────────────────
    import make_figures as mf  # scripts/ is on sys.path

    print("\nGenerating figures ...")
    mf.make_tweets_per_month(df, FIG_TWEETS_MONTH)
    mf.make_predictions_per_account(df, FIG_PREDS_ACCOUNT)
    mf.make_top_tickers(df, FIG_TOP_TICKERS)
    mf.make_accuracy_by_year(df, FIG_ACCURACY_YEAR)
    mf.make_excess_return_hist(df, FIG_EXCESS_RET)
    mf.make_multiple_testing_bar(df, FIG_MULTTEST)
    print("  Calibration curve (RQ3) ...")
    mf.make_calibration_curve(df, FIG_CALIB_CURVE)
    print("  [ROC curves — may take ~2 minutes] ...")
    mf.make_roc_curves(df, FIG_ROC)
    mf.make_confidence_hist(df, FIG_CONF_HIST)
    print("  Regime scatter (RQ4) ...")
    mf.make_regime_scatter(df, FIG_REGIME)

    # ── compute stats for tables ──────────────────────────────────────────────
    print("\nComputing RQ1/RQ2/RQ3/RQ4 stats for tables ...")
    rq1    = compute_rq1(df)
    rq2    = compute_rq2(df)
    calib  = calibration_table(df)
    regime = compute_regime(df)

    # ── build document from scratch ───────────────────────────────────────────
    print("\nBuilding report.docx ...")
    REPORT_PATH.unlink(missing_ok=True)
    doc = new_document()

    write_cover_page(doc, author="<TON NOM>", date="June 2026")

    write_phase0(doc)
    write_phase1(doc)
    write_phase2(doc)
    write_phase3(doc)
    write_phase4(doc, rq1)
    write_phase5(doc, rq2)
    write_phase6(doc, calib)
    write_phase7(doc, regime)
    write_limitations(doc)
    write_conclusion(doc)

    doc.save(str(REPORT_PATH))
    print(f"\nReport saved: {REPORT_PATH}")

    # ── count images and tables ───────────────────────────────────────────────
    from docx import Document as _Doc
    from docx.oxml.ns import qn as _qn
    _d = _Doc(str(REPORT_PATH))
    n_tables = len(_d.tables)
    n_images = sum(
        1
        for para in _d.paragraphs
        for run in para.runs
        for _ in run._r.findall(f".//{_qn('a:blip')}")
    )
    print(f"\nDocument contains {n_images} embedded image(s) and {n_tables} table(s).")


if __name__ == "__main__":
    main()
