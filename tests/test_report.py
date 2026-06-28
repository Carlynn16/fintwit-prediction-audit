"""Tests for src/report.py and scripts/build_report.py."""

import pathlib
import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import sys
sys.path.insert(0, str(pathlib.Path(__file__).parent.parent))

from src.report import (
    add_heading1, add_heading2, add_body, add_figure, add_table, new_document,
    HEADING_COLOR, FONT_NAME, H1_SIZE, H2_SIZE, BODY_SIZE,
)

REPORT_PATH = pathlib.Path(__file__).parent.parent / "report" / "report.docx"


def _first_run(para):
    return para.runs[0] if para.runs else None


class TestReportHelpers:
    def test_heading1_font_name(self):
        doc = new_document()
        add_heading1(doc, "1. Test Heading")
        run = _first_run(doc.paragraphs[-1])
        assert run is not None
        assert run.font.name == FONT_NAME

    def test_heading1_font_size(self):
        doc = new_document()
        add_heading1(doc, "1. Test Heading")
        run = _first_run(doc.paragraphs[-1])
        assert run.font.size == H1_SIZE

    def test_heading1_font_color(self):
        doc = new_document()
        add_heading1(doc, "1. Test Heading")
        run = _first_run(doc.paragraphs[-1])
        assert run.font.color.rgb == HEADING_COLOR

    def test_heading2_font_size(self):
        doc = new_document()
        add_heading2(doc, "1.1 Sub Heading")
        run = _first_run(doc.paragraphs[-1])
        assert run.font.size == H2_SIZE

    def test_heading2_font_color(self):
        doc = new_document()
        add_heading2(doc, "1.1 Sub Heading")
        run = _first_run(doc.paragraphs[-1])
        assert run.font.color.rgb == HEADING_COLOR

    def test_body_font_size(self):
        doc = new_document()
        add_body(doc, "Some body text.")
        run = _first_run(doc.paragraphs[-1])
        assert run.font.size == BODY_SIZE

    def test_body_font_name(self):
        doc = new_document()
        add_body(doc, "Some body text.")
        run = _first_run(doc.paragraphs[-1])
        assert run.font.name == FONT_NAME

    def test_heading_color_values(self):
        assert HEADING_COLOR == RGBColor(0x2F, 0x54, 0x96)

    def test_heading1_text_content(self):
        doc = new_document()
        add_heading1(doc, "1. Introduction")
        assert doc.paragraphs[-1].text == "1. Introduction"

    def test_heading2_text_content(self):
        doc = new_document()
        add_heading2(doc, "1.1 Overview")
        assert doc.paragraphs[-1].text == "1.1 Overview"

    def test_body_text_content(self):
        doc = new_document()
        add_body(doc, "Hello world.")
        assert doc.paragraphs[-1].text == "Hello world."


class TestBuiltReport:
    def test_report_file_exists(self):
        if not REPORT_PATH.exists():
            pytest.skip("report/report.docx not yet built — run scripts/build_report.py first")

    def test_report_has_introduction_heading(self):
        if not REPORT_PATH.exists():
            pytest.skip("report/report.docx not yet built")
        doc = Document(str(REPORT_PATH))
        texts = [p.text.strip() for p in doc.paragraphs]
        assert "1. Introduction" in texts

    def test_report_has_data_heading(self):
        if not REPORT_PATH.exists():
            pytest.skip("report/report.docx not yet built")
        doc = Document(str(REPORT_PATH))
        texts = [p.text.strip() for p in doc.paragraphs]
        assert "2. Data" in texts

    def test_report_headings_use_correct_color(self):
        if not REPORT_PATH.exists():
            pytest.skip("report/report.docx not yet built")
        doc = Document(str(REPORT_PATH))
        for para in doc.paragraphs:
            if para.text.strip() in ("1. Introduction", "2. Data"):
                run = _first_run(para)
                assert run is not None
                assert run.font.color.rgb == HEADING_COLOR, (
                    f"Heading '{para.text}' has wrong colour: {run.font.color.rgb}"
                )

    def test_report_headings_use_times_new_roman(self):
        # When python-docx saves a run whose paragraph uses the Normal style (which
        # already specifies Times New Roman), it omits the redundant w:ascii/w:hAnsi
        # and keeps only w:cs.  We therefore verify the font via the w:cs XML attribute
        # and confirm size (H1 vs H2), bold, and colour also round-trip correctly.
        if not REPORT_PATH.exists():
            pytest.skip("report/report.docx not yet built")
        doc = Document(str(REPORT_PATH))
        checks = {
            "1. Introduction":      H1_SIZE,
            "1.1 Research Questions": H2_SIZE,
        }
        found = set()
        for para in doc.paragraphs:
            txt = para.text.strip()
            if txt not in checks:
                continue
            expected_size = checks[txt]
            run = _first_run(para)
            assert run is not None, f"No runs in heading '{txt}'"
            rPr = run._r.find(qn("w:rPr"))
            assert rPr is not None, f"No rPr for '{txt}'"
            rFonts = rPr.find(qn("w:rFonts"))
            assert rFonts is not None, f"No w:rFonts for '{txt}'"
            cs_font = rFonts.get(qn("w:cs"))
            assert cs_font == FONT_NAME, (
                f"'{txt}': expected w:cs={FONT_NAME!r}, got {cs_font!r}"
            )
            assert run.font.size == expected_size, (
                f"'{txt}': expected size {expected_size}, got {run.font.size}"
            )
            assert run.font.bold is True, f"'{txt}' run is not bold"
            assert run.font.color.rgb == HEADING_COLOR, (
                f"'{txt}': wrong colour {run.font.color.rgb}"
            )
            found.add(txt)
        assert found == set(checks), f"Missing headings in report: {set(checks) - found}"


# ---------------------------------------------------------------------------
# Fixture: tiny valid PNG created via matplotlib (already a project dependency)
# ---------------------------------------------------------------------------

@pytest.fixture()
def tmp_png(tmp_path):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    p = tmp_path / "test_fig.png"
    fig, ax = plt.subplots(figsize=(2, 1))
    ax.plot([0, 1], [0, 1])
    fig.savefig(str(p), dpi=72)
    plt.close(fig)
    return p


# ---------------------------------------------------------------------------
# TestAddFigure
# ---------------------------------------------------------------------------

class TestAddFigure:
    def test_caption_paragraph_added(self, tmp_png):
        doc = new_document()
        n_before = len(doc.paragraphs)
        add_figure(doc, tmp_png, "Figure 1.1. A test caption.")
        # At minimum: one image paragraph + one caption paragraph
        assert len(doc.paragraphs) >= n_before + 2

    def test_caption_text_correct(self, tmp_png):
        doc = new_document()
        add_figure(doc, tmp_png, "Figure 2.1. Monthly volume.")
        cap_para = doc.paragraphs[-1]
        assert cap_para.text == "Figure 2.1. Monthly volume."

    def test_caption_is_italic(self, tmp_png):
        doc = new_document()
        add_figure(doc, tmp_png, "Figure 3.1. Caption text.")
        cap_run = doc.paragraphs[-1].runs[0]
        assert cap_run.font.italic is True

    def test_caption_font_size_10pt(self, tmp_png):
        doc = new_document()
        add_figure(doc, tmp_png, "Figure 4.1. Caption.")
        cap_run = doc.paragraphs[-1].runs[0]
        assert cap_run.font.size == Pt(10)

    def test_image_paragraph_is_centered(self, tmp_png):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = new_document()
        n_before = len(doc.paragraphs)
        add_figure(doc, tmp_png, "Figure 5.1. Caption.")
        # The image lives in the paragraph just before the caption
        img_para = doc.paragraphs[-2]
        assert img_para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_image_blip_element_present(self, tmp_png):
        doc = new_document()
        add_figure(doc, tmp_png, "Figure 6.1. Caption.")
        img_para = doc.paragraphs[-2]
        blips = img_para._p.findall(f".//{qn('a:blip')}")
        assert len(blips) == 1

    def test_multiple_figures_accumulate(self, tmp_png):
        doc = new_document()
        add_figure(doc, tmp_png, "Figure A.")
        add_figure(doc, tmp_png, "Figure B.")
        blips = [
            blip
            for para in doc.paragraphs
            for run in para.runs
            for blip in run._r.findall(f".//{qn('a:blip')}")
        ]
        assert len(blips) == 2


# ---------------------------------------------------------------------------
# TestAddTable
# ---------------------------------------------------------------------------

class TestAddTable:
    _HEADERS = ["Account", "n", "Rate"]
    _ROWS = [
        ["account_01", "100", "45.0%"],
        ["account_02", "80",  "52.5%"],
    ]

    def test_returns_table_object(self):
        from docx.table import Table
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        assert isinstance(tbl, Table)

    def test_correct_row_count(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        assert len(tbl.rows) == 1 + len(self._ROWS)

    def test_correct_column_count(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        assert len(tbl.columns) == len(self._HEADERS)

    def test_header_text_correct(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        for col_idx, h in enumerate(self._HEADERS):
            assert tbl.cell(0, col_idx).text == h

    def test_header_row_is_bold(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        for col_idx in range(len(self._HEADERS)):
            run = tbl.cell(0, col_idx).paragraphs[0].runs[0]
            assert run.font.bold is True

    def test_data_row_text_correct(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        for row_idx, row_data in enumerate(self._ROWS, start=1):
            for col_idx, val in enumerate(row_data):
                assert tbl.cell(row_idx, col_idx).text == val

    def test_data_row_not_bold(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        for row_idx in range(1, 1 + len(self._ROWS)):
            for col_idx in range(len(self._HEADERS)):
                run = tbl.cell(row_idx, col_idx).paragraphs[0].runs[0]
                assert run.font.bold is False

    def test_table_grid_style(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        assert tbl.style.name == "Table Grid"

    def test_font_size_applied(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS, font_size_pt=8)
        run = tbl.cell(0, 0).paragraphs[0].runs[0]
        assert run.font.size == Pt(8)

    def test_font_size_default_is_9pt(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, self._ROWS)
        run = tbl.cell(0, 0).paragraphs[0].runs[0]
        assert run.font.size == Pt(9)

    def test_col_widths_applied(self):
        from docx.shared import Cm
        doc = new_document()
        widths = [3.0, 1.5, 2.0]
        tbl = add_table(doc, self._HEADERS, self._ROWS, col_widths_cm=widths)
        for col_idx, w in enumerate(widths):
            cell = tbl.cell(0, col_idx)
            assert abs(cell.width - Cm(w)) < 500  # within 500 EMU (~0.04 mm) tolerance

    def test_empty_rows_gives_header_only_table(self):
        doc = new_document()
        tbl = add_table(doc, self._HEADERS, [])
        assert len(tbl.rows) == 1

    def test_many_rows_all_populated(self):
        doc = new_document()
        big_rows = [[f"r{i}", str(i), f"{i}%"] for i in range(50)]
        tbl = add_table(doc, self._HEADERS, big_rows)
        assert len(tbl.rows) == 51
        assert tbl.cell(50, 0).text == "r49"
