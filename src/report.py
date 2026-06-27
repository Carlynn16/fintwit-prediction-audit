"""Reusable python-docx helpers that enforce the project report formatting spec.

Spec:
  Font: Times New Roman everywhere
  Body text: 12 pt
  Level-1 headings (1., 2., ...): 18 pt, colour #2F5496
  Sub-headings (2.1, 2.3, ...): 15 pt, colour #2F5496
"""

from __future__ import annotations

import pathlib
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

HEADING_COLOR = RGBColor(0x2F, 0x54, 0x96)
FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
H1_SIZE = Pt(18)
H2_SIZE = Pt(15)


def _set_run_font(run, size: "Pt", bold: bool = False, color: "RGBColor | None" = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Also set the theme font via XML so Word doesn't override with theme defaults
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def add_heading1(doc: Document, text: str) -> None:
    """Add a level-1 heading (18 pt, #2F5496, Times New Roman)."""
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    run = para.add_run(text)
    _set_run_font(run, H1_SIZE, bold=True, color=HEADING_COLOR)


def add_heading2(doc: Document, text: str) -> None:
    """Add a level-2 sub-heading (15 pt, #2F5496, Times New Roman)."""
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    run = para.add_run(text)
    _set_run_font(run, H2_SIZE, bold=True, color=HEADING_COLOR)


def add_body(doc: Document, text: str) -> None:
    """Add a body paragraph (12 pt, Times New Roman, black)."""
    para = doc.add_paragraph()
    para.style = doc.styles["Normal"]
    run = para.add_run(text)
    _set_run_font(run, BODY_SIZE)


def new_document() -> Document:
    """Return a blank Document with default Normal style set to Times New Roman 12pt."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    return doc


def open_or_create(path: pathlib.Path) -> Document:
    """Open an existing report .docx or create a new one if it doesn't exist."""
    if path.exists():
        return Document(str(path))
    return new_document()


def add_figure(doc: Document, path, caption: str) -> None:
    """Add a centered image (6.5 in wide) followed by a centered italic caption."""
    pic_para = doc.add_picture(str(path), width=Inches(6.5))
    # add_picture returns the paragraph that contains the image
    pic_para = doc.paragraphs[-1]
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_para.add_run(caption)
    _set_run_font(run, Pt(10))
    run.font.italic = True


def add_table(
    doc: Document,
    headers: list,
    rows: list,
    col_widths_cm: list | None = None,
    font_size_pt: int = 9,
):
    """Create a Word table with header row and data rows.

    Parameters
    ----------
    doc           : Document to add the table to
    headers       : list of column header strings
    rows          : list of row iterables (one per data row)
    col_widths_cm : optional list of column widths in cm
    font_size_pt  : font size for all cells (default 9 pt)

    Returns the created Table object.
    """
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # Header row (index 0)
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(header_text))
        _set_run_font(run, Pt(font_size_pt), bold=True)

    # Data rows (indices 1+)
    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(cell_text))
            _set_run_font(run, Pt(font_size_pt), bold=False)

    # Apply column widths if provided
    if col_widths_cm is not None:
        for col_idx, w in enumerate(col_widths_cm):
            for row_idx in range(n_rows):
                table.cell(row_idx, col_idx).width = Cm(w)

    return table


def write_cover_page(
    doc: Document,
    author: str = "<TON NOM>",
    date: str = "June 2026",
) -> None:
    """Insert a centred cover page followed by a hard page break."""
    from docx.enum.text import WD_BREAK

    def _c(text: str, size_pt: float, bold: bool = False, italic: bool = False, color=None):
        para = doc.add_paragraph()
        para.style = doc.styles["Normal"]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            run = para.add_run(text)
            _set_run_font(run, Pt(size_pt), bold=bold, color=color)
            run.font.italic = italic
        return para

    # top padding (~1/3 page)
    for _ in range(10):
        _c("", 12)

    _c("FinTwit Prediction Audit", 28, bold=True, color=HEADING_COLOR)

    _c("", 12)

    _c(
        "An LLM-Powered Statistical Audit of ~18,000 Stock Predictions",
        16, color=HEADING_COLOR,
    )
    _c(
        "by 51 Financial Influencers on X (2021–2025)",
        16, color=HEADING_COLOR,
    )

    _c("", 12)
    _c("", 12)

    _c(
        "Aggregate directional accuracy 45.0%; "
        "no account shows durable skill across market regimes",
        12, italic=True,
    )

    for _ in range(8):
        _c("", 12)

    _c(author, 12, bold=True)
    _c(date, 12)

    # hard page break
    br_para = doc.add_paragraph()
    br_para.style = doc.styles["Normal"]
    br_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = br_para.add_run()
    run.add_break(WD_BREAK.PAGE)
